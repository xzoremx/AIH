import fitz  # PyMuPDF
import os, sys, re, tkinter as tk
from tkinter import messagebox
from docx import Document

# ================== base_dir ==================
if hasattr(sys, '_MEIPASS'):
    base_dir = sys._MEIPASS
else:
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

def preguntar_inacal():
    try:
        root = tk.Tk(); root.withdraw()
        r = messagebox.askyesno("Certificación INACAL", "¿El ensayo está certificado por INACAL?")
        root.destroy()
        return r
    except Exception:
        return bool(int(os.environ.get("INACAL", "0")))

esta_certificado = preguntar_inacal()
pdf_encabezado_path = os.path.join(
    base_dir, "Formatos_Página",
    "formato1_superior_con_sello.pdf" if esta_certificado else "formato1_superior_sin_sello.pdf"
)

sys.stdout.reconfigure(encoding='utf-8')

# Rutas
carpeta_pdfs_tablas = os.path.join(base_dir, "PDFs_Impresos")
pdf_pie_path = os.path.join(base_dir, "Formatos_Página", "formato1_inferior.pdf")
carpeta_salida = os.path.join(base_dir, "PDFs_Finales")
carpeta_word = base_dir

def extraer_encabezado_word(carpeta_word):
    docx_files = [f for f in os.listdir(carpeta_word) if f.lower().endswith('.docx')]
    if not docx_files:
        print("No se encontraron archivos .docx en la carpeta."); return ""
    word_path = os.path.join(carpeta_word, docx_files[0])
    try:
        doc = Document(word_path)
    except Exception as e:
        print("No se pudo abrir el .docx:", e); return ""
    patron = re.compile(r'INFORME\s+DE\s+ENSAYO\s+(?:N[°ºo]\.?)[\s]*\d{1,7}-\d{3,6}/\d{2,4}', re.IGNORECASE)
    for para in doc.paragraphs:
        m = patron.search(para.text)
        if m: return m.group()
    try:
        for para in doc.sections[0].header.paragraphs:
            m = patron.search(para.text)
            if m: return m.group()
    except Exception as e:
        print("No se pudo acceder al encabezado del documento:", e)
    print("No se encontró el encabezado en el archivo Word.")
    return ""

encabezado_extraido = extraer_encabezado_word(carpeta_word)
print("Encabezado extraído:", encabezado_extraido or "(vacío)")

# ================== Parámetros ==================
altura_encabezado = 100    # pt
altura_pie_pagina = 100     # pt
ancho_lateral_izquierdo = 30  # pt (nuevo)
ancho_lateral_derecho  = 30  # pt (antes: 30)

texto_lateral = '" EL USO INDEBIDO DE ESTE INFORME DE ENSAYO CONSTITUYE DELITO SANCIONADO CONFORME A LA LEY, POR LA AUTORIDAD COMPETENTE "'

# Validaciones
os.makedirs(carpeta_salida, exist_ok=True)
for path in (pdf_encabezado_path, pdf_pie_path, carpeta_pdfs_tablas):
    if not os.path.exists(path):
        raise FileNotFoundError(f"Falta recurso: {path}")

# ================== Composición (lienzo expandido simétrico) ==================
with fitz.open(pdf_encabezado_path) as pdf_encabezado, fitz.open(pdf_pie_path) as pdf_pie:
    w_head_src, h_head_src = pdf_encabezado[0].rect.width, pdf_encabezado[0].rect.height
    w_pie_src,  h_pie_src  = pdf_pie[0].rect.width,        pdf_pie[0].rect.height

    archivos = sorted([f for f in os.listdir(carpeta_pdfs_tablas) if f.lower().endswith(".pdf")])
    if not archivos:
        print("No hay PDFs en 'PDFs_Impresos'.")
    for idx, nombre_archivo in enumerate(archivos, 1):
        ruta_tabla = os.path.join(carpeta_pdfs_tablas, nombre_archivo)
        with fitz.open(ruta_tabla) as pdf_tabla, fitz.open() as pdf_final:
            print(f"[{idx}/{len(archivos)}] Procesando: {nombre_archivo}")
            for pagina_tabla in pdf_tabla:
                w_tab, h_tab = pagina_tabla.rect.width, pagina_tabla.rect.height

                # Lienzo final: sumamos margen IZQ y DER
                w_util = max(w_tab, w_head_src, w_pie_src)
                w_final = w_util + ancho_lateral_izquierdo + ancho_lateral_derecho
                h_final = h_tab + altura_encabezado + altura_pie_pagina

                page = pdf_final.new_page(width=w_final, height=h_final)

                # Franja de encabezado (entre márgenes laterales)
                rect_head_target = fitz.Rect(
                    ancho_lateral_izquierdo, 0,
                    w_final - ancho_lateral_derecho, altura_encabezado
                )
                page.show_pdf_page(rect_head_target, pdf_encabezado, 0)

                # Texto dinámico del Word (centrado bajo el encabezado, respetando márgenes)
                if encabezado_extraido:
                    page.insert_textbox(
                        fitz.Rect(
                            ancho_lateral_izquierdo + 12,
                            altura_encabezado - 22,
                            w_final - ancho_lateral_derecho - 12,
                            altura_encabezado + 28
                        ),
                        encabezado_extraido, fontsize=11, fontname="Helvetica-Bold", align=1, color=(0,0,0)
                    )

                # Tabla SIN ESCALADO: colocada debajo del encabezado, desplazada por el margen izquierdo
                rect_tabla_destino = fitz.Rect(
                    ancho_lateral_izquierdo,
                    altura_encabezado,
                    ancho_lateral_izquierdo + w_tab,
                    altura_encabezado + h_tab
                )
                page.show_pdf_page(rect_tabla_destino, pdf_tabla, pagina_tabla.number)

                # Franja de pie (entre márgenes laterales)
                rect_pie_target = fitz.Rect(
                    ancho_lateral_izquierdo,
                    h_final - altura_pie_pagina,
                    w_final - ancho_lateral_derecho,
                    h_final
                )
                page.show_pdf_page(rect_pie_target, pdf_pie, 0)

                # Banda lateral DERECHA con el texto rotado
                page.insert_textbox(
                    fitz.Rect(w_final - ancho_lateral_derecho + 8, 0, w_final, h_final),
                    texto_lateral, fontsize=7, rotate=90, fontname="Helvetica-Bold", align=1, color=(0,0,0)
                )

                # (Opcional) Banda lateral IZQUIERDA vacía para simetría visual
                # Si deseas poner algo aquí, usa otro insert_textbox sobre:
                # fitz.Rect(0, 0, ancho_lateral_izquierdo, h_final)

            # Guardar
            nombre_final = os.path.splitext(nombre_archivo)[0] + "_final.pdf"
            ruta_salida = os.path.join(carpeta_salida, nombre_final)
            pdf_final.save(ruta_salida)
            print(f" PDF combinado: {nombre_final}")

print(" Todos los PDFs finales han sido generados correctamente.")


